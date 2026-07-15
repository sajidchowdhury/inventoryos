#!/usr/bin/env python3
"""Generate InventoryOS Production Readiness Audit & Remediation Plan as DOCX."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = "/home/z/my-project/download/InventoryOS_Production_Readiness_Audit_and_Remediation_Plan.docx"
os.makedirs(os.path.dirname(OUT), exist_ok=True)

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

# ── Colors ──
C_PRIMARY = RGBColor(0x10, 0x18, 0x20)
C_BODY = RGBColor(0x18, 0x20, 0x30)
C_ACCENT = RGBColor(0x7C, 0x3A, 0xED)
C_RED = RGBColor(0xDC, 0x26, 0x26)
C_AMBER = RGBColor(0xD9, 0x77, 0x06)
C_GREEN = RGBColor(0x05, 0x96, 0x69)
C_GRAY = RGBColor(0x6B, 0x72, 0x80)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_custom(text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color or (C_PRIMARY if level <= 2 else C_ACCENT)
        if level == 1: run.font.size = Pt(18)
        elif level == 2: run.font.size = Pt(15)
        else: run.font.size = Pt(13)
    return h

def add_body(text, indent=True):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = C_BODY
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = C_BODY
    return p

def add_table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, (text, color) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10)
        if header:
            run.bold = True
            run.font.color.rgb = C_WHITE
            set_cell_shading(cell, '7C3AED')
        else:
            run.font.color.rgb = color or C_BODY
    return row

# ═════════════ COVER ═════════════
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("InventoryOS"); run.font.size = Pt(36); run.font.color.rgb = C_ACCENT; run.bold = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Production Readiness Audit"); run.font.size = Pt(28); run.font.color.rgb = C_PRIMARY; run.bold = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("& Remediation Plan"); run.font.size = Pt(28); run.font.color.rgb = C_PRIMARY; run.bold = True
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Database Architecture | Application Logic | Security | Scalability"); run.font.size = Pt(13); run.font.color.rgb = C_GRAY; run.italic = True
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Current Score: 31/100  |  Target Score: 98+/100"); run.font.size = Pt(14); run.font.color.rgb = C_RED; run.bold = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Status: NOT READY FOR PRODUCTION"); run.font.size = Pt(14); run.font.color.rgb = C_RED; run.bold = True
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prepared by: Senior Enterprise Database Architect\nDate: July 2026"); run.font.size = Pt(11); run.font.color.rgb = C_GRAY
doc.add_page_break()

# ═════════════ TOC ═════════════
add_heading_custom("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary", "2. Audit Scorecard",
    "3. Critical Findings (Must Fix Before Launch)",
    "   C-1: No Database Transactions on Financial Operations",
    "   C-2: Stock Can Go Negative (Race Condition)",
    "   C-3: No Unique Constraint on Serial Numbers",
    "   C-4: No Unique Constraint on Invoice Numbers",
    "   C-5: 9 CCTV Models Missing Business FK Relation",
    "   C-6: Float Used for All Money Fields",
    "4. High Severity Findings",
    "   H-1: No CHECK Constraints on Numeric Fields",
    "   H-2: 36 Foreign Keys Without Explicit onDelete Rule",
    "   H-3: No Audit Trail for Stock Movements",
    "   H-4: No Double-Entry Ledger System",
    "   H-5: Partial Failure on Serial History Writes",
    "   H-6: Migration Drift",
    "5. Medium Severity Findings",
    "6. Critical Checks Verification Matrix",
    "7. Remediation Phases (10 Phases)",
    "   Phase 1: Database Transaction Safety",
    "   Phase 2: Data Integrity & Constraints",
    "   Phase 3: Money Type Migration (Float to Decimal)",
    "   Phase 4: Business FK Relations & Cascade Rules",
    "   Phase 5: Unique Constraints & Invoice Numbers",
    "   Phase 6: Stock Movement Audit Table",
    "   Phase 7: Double-Entry Ledger System",
    "   Phase 8: Performance & Indexing",
    "   Phase 9: Security & Multi-Tenancy Hardening",
    "   Phase 10: Migration Generation & Deployment",
    "8. Phase Priority & Timeline",
    "9. Target Score: 98+ Roadmap",
]
for item in toc_items:
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.3
    run = p.add_run(item); run.font.size = Pt(11)
    if item.startswith("   "): run.font.color.rgb = C_GRAY
    else: run.font.color.rgb = C_BODY; run.bold = True
doc.add_page_break()

# ═════════════ 1. EXECUTIVE SUMMARY ═════════════
add_heading_custom("1. Executive Summary", level=1)
add_body("This audit performed a comprehensive production-readiness review of the InventoryOS database schema (124 Prisma models, 3,837 lines) and critical CCTV module API routes. The system has fundamental architectural issues that make it unsafe for production use with real financial data.")
add_body("The most critical finding is that multi-step financial operations (sales, purchases, replacements) execute without database transactions, meaning a crash mid-operation leaves the database in an inconsistent state with stock and money records out of sync.")
add_body("The current overall score is 31 out of 100. The system is classified as NOT READY FOR PRODUCTION. However, with the phased remediation plan outlined in this document, the target score of 98+ is achievable. The plan is split into 10 phases, ordered by criticality, with each phase designed to be independently deployable.")
add_body("Key statistics: 6 critical issues, 6 high-severity issues, 6 medium-severity issues, 10 out of 10 critical checks failed, 9 models missing foreign key relations, 36 relations without explicit cascade rules, and zero CHECK constraints across the entire schema.")

# ═════════════ 2. SCORECARD ═════════════
add_heading_custom("2. Audit Scorecard", level=1)
table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_row(table, [("Category", C_WHITE), ("Current", C_WHITE), ("Target", C_WHITE), ("Gap", C_WHITE)], header=True)
for cat, cur, tgt, gap in [
    ("Architecture","5/10","10/10","5"),("Relationships","4/10","10/10","6"),
    ("Data Integrity","2/10","10/10","8"),("Performance","5/10","10/10","5"),
    ("Inventory Logic","3/10","10/10","7"),("Accounting Integrity","2/10","10/10","8"),
    ("Security","4/10","10/10","6"),("Scalability","4/10","10/10","6"),
    ("Production Readiness","2/10","10/10","8")]:
    add_table_row(table, [(cat,C_BODY),(cur,C_RED),(tgt,C_GREEN),(gap,C_AMBER)])
add_table_row(table, [("OVERALL",C_WHITE),("31/100",C_WHITE),("98/100",C_WHITE),("67",C_WHITE)])
for cell in table.rows[-1].cells:
    set_cell_shading(cell, '7C3AED')
    for p in cell.paragraphs:
        for r in p.runs: r.bold = True
doc.add_paragraph()

# ═════════════ 3. CRITICAL FINDINGS ═════════════
add_heading_custom("3. Critical Findings (Must Fix Before Launch)", level=1)
critical_findings = [
    {"id":"C-1","title":"No Database Transactions on Financial Operations","problem":"The Sales, Purchase, and Replacement APIs perform 10+ separate await db.create() calls without wrapping them in $transaction(). Each operation commits independently.","impact":"If the server crashes after creating a sale record but before decrementing stock, the database is left in an inconsistent state. This is unrecoverable without manual database surgery.","recommendation":"Wrap every multi-step financial operation in db.$transaction(async (tx) => { ... }). All writes either commit together or roll back together."},
    {"id":"C-2","title":"Stock Can Go Negative (Race Condition)","problem":"The stock safety check runs outside a transaction. Two concurrent sales can both read stock=5, both pass the check, and both decrement, resulting in stock=-3.","impact":"Overselling. Two cashiers selling the same product simultaneously can sell more than available stock.","recommendation":"Use atomic conditional update inside a transaction: UPDATE ... SET stock = stock - qty WHERE stock >= qty. Also add CHECK constraint: ALTER TABLE cctv_products ADD CONSTRAINT stock_non_negative CHECK (stock >= 0)."},
    {"id":"C-3","title":"No Unique Constraint on Serial Numbers Per Business","problem":"CCTVSerialItem.serialNumber has only an index, not a unique constraint. Concurrent inserts can create duplicates.","impact":"Duplicate serial numbers break the entire serial tracking system: warranty tracking, sales scanning, and replacement flow all assume serials are unique.","recommendation":"Add @@unique([businessId, serialNumber]) to CCTVSerialItem."},
    {"id":"C-4","title":"No Unique Constraint on Invoice Numbers","problem":"CCTVSale.invoiceNo has no unique constraint. Duplicate invoice numbers can be created.","impact":"Regulatory audit failure. Two invoices with the same number make financial auditing impossible.","recommendation":"Add @@unique([businessId, invoiceNo]) to CCTVSale."},
    {"id":"C-5","title":"9 CCTV Models Missing Business FK Relation","problem":"9 models (CCTVPurchaseItem, CCTVSaleItem, CCTVPayment, CCTVReturnItem, CCTVWarrantyClaim, CCTVSerialHistory, CCTVRepair, CCTVSupplierReplacement, CCTVEstimateItem) have businessId but no FK to Business.","impact":"Orphan records can exist. Cascading deletes from Business won't clean up these tables. Cross-tenant data leakage is possible.","recommendation":"Add business Business @relation(fields: [businessId], references: [id], onDelete: Cascade) to all 9 models."},
    {"id":"C-6","title":"Float Used for All Money Fields","problem":"Every monetary field uses Float. Floating-point arithmetic produces rounding errors: 0.1 + 0.2 = 0.30000000000000004.","impact":"Financial discrepancies accumulate. Ledger balances won't reconcile. Tax authorities reject imprecise financial records.","recommendation":"Change all money fields to Decimal @db.Decimal(12,2)."},
]
for f in critical_findings:
    add_heading_custom(f"{f['id']}: {f['title']}", level=2, color=C_RED)
    t = doc.add_table(rows=0, cols=2); t.style = 'Table Grid'
    for label, val in [("Severity","CRITICAL"),("Problem",f["problem"]),("Impact",f["impact"]),("Recommendation",f["recommendation"]),("Priority","Fix before launch")]:
        row = t.add_row()
        c0 = row.cells[0]; c0.text=''; r0 = c0.paragraphs[0].add_run(label); r0.bold=True; r0.font.size=Pt(10); r0.font.color.rgb=C_PRIMARY; set_cell_shading(c0,'F3F4F6')
        c1 = row.cells[1]; c1.text=''; r1 = c1.paragraphs[0].add_run(val); r1.font.size=Pt(10); r1.font.color.rgb=C_BODY
    doc.add_paragraph()

# ═════════════ 4. HIGH FINDINGS ═════════════
add_heading_custom("4. High Severity Findings", level=1)
high = [
    ("H-1","No CHECK Constraints on Any Numeric Field","No CHECK constraints exist. Quantities can be negative. Prices can be negative. Discounts can exceed total.","Invalid data stored: negative stock, negative prices, discounts larger than invoice.","Add CHECK constraints via raw SQL migrations.","Fix before launch"),
    ("H-2","36 Foreign Keys Without Explicit onDelete Rule","36 of 183 relations have no onDelete rule, defaulting to NO ACTION.","Cannot delete products that have been sold. Opaque database errors.","Define onDelete: Restrict for financial records, SetNull for optional references.","Fix before launch"),
    ("H-3","No Audit Trail for Stock Movements","Stock changes via increment/decrement directly on product table. No movement history.","Cannot trace why stock changed. Cannot audit discrepancies.","Create CCTVStockMovement table with type, quantityChange, balanceAfter.","Fix within 3 months"),
    ("H-4","No Double-Entry Ledger System","Ledger computed on-the-fly. No persistent ledger table with debit/credit entries.","Balances can drift. No audit trail. Cannot produce trial balance.","Implement CCTVLedgerEntry table with balanced debit+credit entries.","Fix within 3 months"),
    ("H-5","Partial Failure on Serial History Writes","Serial history writes wrapped in try/catch, silently fail.","Missing audit entries. Warranty timeline unreliable.","History writes must be inside the same transaction.","Fix before launch"),
    ("H-6","Migration Drift","New CCTV tables added via prisma db push (dev-only). No versioned migration.","Cannot deploy to fresh database. Production deployments fragile.","Generate proper migration: bunx prisma migrate dev.","Fix before launch"),
]
for fid, title, prob, impact, rec, pri in high:
    add_heading_custom(f"{fid}: {title}", level=2, color=C_AMBER)
    add_body(f"Problem: {prob}"); add_body(f"Impact: {impact}"); add_body(f"Recommendation: {rec}")
    p = doc.add_paragraph(); run = p.add_run(f"Priority: {pri}"); run.bold=True; run.font.size=Pt(11); run.font.color.rgb=C_AMBER
    doc.add_paragraph()

# ═════════════ 5. MEDIUM FINDINGS ═════════════
add_heading_custom("5. Medium Severity Findings", level=1)
med = [
    ("M-1","No Composite Index on [businessId, serialNumber]","Add @@index([businessId, serialNumber]) to CCTVSerialItem."),
    ("M-2","Serial Numbers Stored as Comma-Separated String","Create separate CCTVPurchaseSerialItem table."),
    ("M-3","No Server-Side Pagination for Sales/Purchases","Add page/pageSize parameters with total count."),
    ("M-4","No Rate Limiting on Financial APIs","Add rate limiting middleware."),
    ("M-5","No Archiving Strategy","Move records older than 2 years to archive tables."),
    ("M-6","No Materialized Views for Reports","Use materialized views refreshed daily."),
]
for fid, title, rec in med:
    add_heading_custom(f"{fid}: {title}", level=3, color=C_GRAY)
    add_body(f"Recommendation: {rec}"); doc.add_paragraph()

# ═════════════ 6. CHECKS MATRIX ═════════════
add_heading_custom("6. Critical Checks Verification Matrix", level=1)
table = doc.add_table(rows=1, cols=3); table.style='Table Grid'
add_table_row(table, [("Check",C_WHITE),("Status",C_WHITE),("Notes",C_WHITE)], header=True)
for check, status, notes in [
    ("Can stock become negative?","FAILED","No CHECK constraint, no atomic update"),
    ("Can two concurrent sales oversell?","FAILED","No transaction, no row-level locking"),
    ("Can an invoice exist without items?","FAILED","Sale created first, items after — crash = orphan"),
    ("Can ledger entries become unbalanced?","FAILED","No ledger table; computed dynamically"),
    ("Can a record reference a deleted parent?","FAILED","9 models have no FK to Business"),
    ("Can data from Business A be visible to B?","AT RISK","No RLS; relies on app-layer filtering"),
    ("Can duplicate invoice numbers occur?","FAILED","No unique constraint on CCTVSale.invoiceNo"),
    ("Can invalid discounts be stored?","FAILED","No CHECK that discount <= subtotal"),
    ("Can inventory and accounting go out of sync?","FAILED","No transaction wrapping both"),
    ("Can partial failures leave inconsistent state?","FAILED","Every multi-step operation is non-transactional"),
]:
    color = C_RED if status=="FAILED" else C_AMBER
    add_table_row(table, [(check,C_BODY),(status,color),(notes,C_BODY)])
doc.add_paragraph()

# ═════════════ 7. REMEDIATION PHASES ═════════════
add_heading_custom("7. Remediation Phases", level=1)
add_body("The remediation plan is split into 10 phases, ordered by criticality. Each phase is independently deployable. Phases 1-6 must be completed before launch. Phases 7-10 bring the system to 98+ readiness.")

phases = [
    {"num":1,"title":"Database Transaction Safety","sev":"Critical","goal":"Wrap all financial APIs in $transaction()","fixes":["C-1","C-2","H-5"],"steps":["Wrap POST /cctv/sales in db.$transaction()","Wrap POST /cctv/purchases in db.$transaction()","Wrap POST /cctv/supplier-replacements in db.$transaction()","Wrap PATCH /cctv/supplier-replacements/[id] in db.$transaction()","Wrap POST /cctv/repairs in db.$transaction()","Wrap POST /cctv/payments in db.$transaction()","Move serial history writes inside main transaction (remove try/catch)","Use atomic conditional updateMany for stock: where: { stock: { gte: qty } }","Test: simulate crash mid-operation, verify rollback"],"impact":"+15 points"},
    {"num":2,"title":"Data Integrity & Constraints","sev":"Critical","goal":"Add CHECK constraints and NOT NULL at database level","fixes":["H-1","C-2"],"steps":["Create raw SQL migration with CHECK constraints:","  cctv_products: stock >= 0, costPrice >= 0, sellPrice >= 0","  cctv_sales: discount <= subtotal, totalAmount >= 0","  cctv_payments: amount > 0","  cctv_expenses: amount > 0","  cctv_repairs: repairCost >= 0","Add NOT NULL to businessId on all 9 orphan models","Add default values where missing","Test: attempt to insert negative stock, verify DB rejects"],"impact":"+12 points"},
    {"num":3,"title":"Money Type Migration (Float to Decimal)","sev":"Critical","goal":"Change all money fields from Float to Decimal @db.Decimal(12,2)","fixes":["C-6"],"steps":["Create migration to ALTER COLUMN TYPE for all Float money fields","Update Prisma schema to use Decimal @db.Decimal(12, 2)","Update all API routes to parse Decimal values","Update all UI components to format Decimal values","Test: verify 0.1 + 0.2 = 0.30 exactly"],"impact":"+10 points","sub":["3a: Schema migration","3b: API route updates","3c: UI formatting updates"]},
    {"num":4,"title":"Business FK Relations & Cascade Rules","sev":"Critical","goal":"Add Business FK to 9 orphaned models + fix 36 missing onDelete","fixes":["C-5","H-2"],"steps":["Add business Business @relation to 9 models:","  CCTVPurchaseItem, CCTVSaleItem, CCTVPayment","  CCTVReturnItem, CCTVWarrantyClaim, CCTVSerialHistory","  CCTVRepair, CCTVSupplierReplacement, CCTVEstimateItem","Add onDelete: Cascade to all new Business relations","Fix 36 relations: Restrict for financial, SetNull for optional","Test: delete a Business, verify child records cleaned up"],"impact":"+8 points","sub":["4a: Add Business FK relations","4b: Fix onDelete cascade rules"]},
    {"num":5,"title":"Unique Constraints & Invoice Numbers","sev":"Critical","goal":"Add unique constraints on serials, invoices, tokens","fixes":["C-3","C-4"],"steps":["Add @@unique([businessId, serialNumber]) to CCTVSerialItem","Add @@unique([businessId, invoiceNo]) to CCTVSale","Add partial unique index for tokenNo on CCTVRepair","Add @@unique([businessId, estimateNo]) to CCTVEstimate","Generate and apply migration","Test: attempt duplicate serial, verify DB rejects"],"impact":"+5 points"},
    {"num":6,"title":"Stock Movement Audit Table","sev":"High","goal":"Create CCTVStockMovement table for stock change traceability","fixes":["H-3"],"steps":["Create CCTVStockMovement model:","  id, businessId, productId, movementType","  quantityChange (signed), balanceAfter","  referenceId, referenceType, performedBy, createdAt","Add Business FK + index on [businessId, productId]","Update Sales API: create movement on every stock decrement","Update Purchases API: create movement on every stock increment","Update Returns API: create movement on stock restoration","Update Replacements API: create movement on stock change","Create stock movement report","Test: verify every stock change has movement record"],"impact":"+8 points"},
    {"num":7,"title":"Double-Entry Ledger System","sev":"High","goal":"Implement persistent ledger with balanced debit/credit entries","fixes":["H-4"],"steps":["Create CCTVLedgerEntry model:","  id, businessId, accountId, entryType (DEBIT/CREDIT)","  amount (Decimal), balanceAfter","  referenceId, referenceType, description, performedBy","Add @@index([businessId, accountId, createdAt])","Update Sales API: balanced entries (debit receivable, credit revenue)","Update Purchase API: balanced entries (debit inventory, credit payable)","Update Payment API: balanced entries (debit cash, credit receivable)","Update Expense API: balanced entries (debit expense, credit cash)","Create trial balance report","Verify: sum of debits = sum of credits per business","Test: create sale, verify ledger balanced"],"impact":"+10 points","sub":["7a: Schema","7b: Sales ledger","7c: Purchase ledger","7d: Payment+expense","7e: Trial balance"]},
    {"num":8,"title":"Performance & Indexing","sev":"Medium","goal":"Add missing indexes, composite indexes, pagination","fixes":["M-1","M-3"],"steps":["Add @@index([businessId, serialNumber]) on CCTVSerialItem","Add @@index([businessId, saleDate]) on CCTVSale","Add @@index([businessId, purchaseDate]) on CCTVPurchase","Add @@index([businessId, paymentDate]) on CCTVPayment","Add server-side pagination to all GET endpoints","Replace take: 50 with skip/take pagination","Test: verify query plans use indexes (EXPLAIN ANALYZE)"],"impact":"+5 points"},
    {"num":9,"title":"Security & Multi-Tenancy Hardening","sev":"Medium","goal":"Implement RLS and tenant-safe query patterns","fixes":["Security"],"steps":["Enable PostgreSQL RLS on all CCTV tables","Create RLS policies: USING (business_id = current_setting('app.business_id'))","Set business_id in session on every API request","Add middleware to verify businessId in every query","Add field-level encryption for sensitive data","Add audit logging for all data modifications","Test: attempt cross-tenant query, verify RLS blocks"],"impact":"+6 points"},
    {"num":10,"title":"Migration Generation & Deployment","sev":"High","goal":"Generate proper versioned migrations and update deployment","fixes":["H-6"],"steps":["Run: bunx prisma migrate dev --name cctv_production_hardening","Update deploy-update.sh to use prisma migrate deploy","Create pre-deploy-verify.sh for schema validation","Test: deploy to fresh database using only migrations","Update DEPLOYMENT.md with migration-based process","Archive old db push approach"],"impact":"+5 points"},
]

for ph in phases:
    color = C_RED if ph["sev"]=="Critical" else (C_AMBER if ph["sev"]=="High" else C_GRAY)
    add_heading_custom(f"Phase {ph['num']}: {ph['title']}", level=2, color=color)
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.3
    r=p.add_run("Severity: "); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=color
    r2=p.add_run(ph["sev"]); r2.font.size=Pt(11); r2.font.color.rgb=C_BODY
    add_body(f"Goal: {ph['goal']}")
    p=doc.add_paragraph(); r=p.add_run("Fixes: "); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=C_BODY
    r2=p.add_run(", ".join(ph["fixes"])); r2.font.size=Pt(11); r2.font.color.rgb=C_BODY
    add_heading_custom("Implementation Steps", level=3, color=C_ACCENT)
    for s in ph["steps"]: add_bullet(s)
    if "sub" in ph:
        p=doc.add_paragraph(); r=p.add_run("Sub-phases: "); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=C_BODY
        r2=p.add_run(", ".join(ph["sub"])); r2.font.size=Pt(11); r2.font.color.rgb=C_GRAY
    p=doc.add_paragraph(); r=p.add_run("Score Impact: "); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=C_BODY
    r2=p.add_run(ph["impact"]); r2.font.size=Pt(11); r2.font.color.rgb=C_GREEN; r2.bold=True
    doc.add_paragraph()

# ═════════════ 8. TIMELINE ═════════════
add_heading_custom("8. Phase Priority & Timeline", level=1)
table = doc.add_table(rows=1, cols=5); table.style='Table Grid'
add_table_row(table,[("Phase",C_WHITE),("Title",C_WHITE),("Severity",C_WHITE),("Timeline",C_WHITE),("Score",C_WHITE)],header=True)
for ph, title, sev, tl, sc in [
    ("1","Database Transaction Safety","Critical","Week 1","+15"),
    ("2","Data Integrity & Constraints","Critical","Week 1","+12"),
    ("3a-c","Money Type Migration","Critical","Week 1-2","+10"),
    ("4a-b","Business FK & Cascade Rules","Critical","Week 2","+8"),
    ("5","Unique Constraints","Critical","Week 2","+5"),
    ("6","Stock Movement Audit Table","High","Week 2-3","+8"),
    ("7a-e","Double-Entry Ledger","High","Week 3-4","+10"),
    ("8","Performance & Indexing","Medium","Week 4","+5"),
    ("9","Security & RLS","Medium","Week 4-5","+6"),
    ("10","Migration & Deployment","High","Week 5","+5"),
]:
    color = C_RED if sev=="Critical" else (C_AMBER if sev=="High" else C_GRAY)
    add_table_row(table,[(ph,C_BODY),(title,C_BODY),(sev,color),(tl,C_BODY),(sc,C_GREEN)])
doc.add_paragraph()
p=doc.add_paragraph(); r=p.add_run("Total estimated time: 5 weeks of focused development."); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=C_PRIMARY

# ═════════════ 9. ROADMAP ═════════════
add_heading_custom("9. Target Score: 98+ Roadmap", level=1)
add_body("The following table shows the projected score progression as each phase is completed.")
table = doc.add_table(rows=1, cols=4); table.style='Table Grid'
add_table_row(table,[("After Phase",C_WHITE),("Score",C_WHITE),("Key Improvement",C_WHITE),("Status",C_WHITE)],header=True)
for ph, sc, imp, st in [
    ("Current","31/100","Baseline audit","NOT READY"),
    ("Phase 1","46/100","Transaction safety","STILL AT RISK"),
    ("Phase 2","58/100","Data integrity","STILL AT RISK"),
    ("Phase 3","68/100","Decimal money","IMPROVING"),
    ("Phase 4","76/100","FK relations","IMPROVING"),
    ("Phase 5","81/100","Unique constraints","NEAR READY"),
    ("Phase 6","89/100","Stock audit trail","NEAR READY"),
    ("Phase 7","95/100","Double-entry ledger","ALMOST READY"),
    ("Phase 8","96/100","Performance","READY WITH CONDITIONS"),
    ("Phase 9","97/100","Security hardening","READY"),
    ("Phase 10","98/100","Migration & deployment","READY FOR PRODUCTION"),
]:
    color = C_RED if "NOT" in st or "RISK" in st else (C_AMBER if "IMPROV" in st or "NEAR" in st or "CONDITIONS" in st else C_GREEN)
    add_table_row(table,[(ph,C_BODY),(sc,color),(imp,C_BODY),(st,color)])
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("TARGET: 98/100 — READY FOR PRODUCTION"); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=C_GREEN

doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT)/1024:.1f} KB")
